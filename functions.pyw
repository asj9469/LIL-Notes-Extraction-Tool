import re
import tkinter
from tkinter import messagebox
from docx import Document
from docx.text.run import Font
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_LINE_SPACING

#functions
def set_input():
    main_UI.input_path= main_UI.filedialog.askopenfilename(filetypes = [("Text Document",".txt")])
    if main_UI.input_path: # user selected file 
        main_UI.input_button.configure(text= "file selected",  font=main_UI.custom_font)
                
def set_output():
    main_UI.output_path = main_UI.filedialog.asksaveasfilename(filetypes=[("Text Document",".txt"), ("Markdown File", ".md"),("Microsoft Word Document",".docx")])
    if main_UI.output_path: # user selected path 
        main_UI.output_button.configure(text="path selected", font=main_UI.custom_font)

def extract_event():
    v = main_UI.optionmenu_var
    i = main_UI.input_path
    o = main_UI.output_path

    if(v.get()==".txt"): #if user selected .txt file
        if (o[-4:])!= ".txt": #if the last 4 characters are not the txt extension, concat the extension to path
            o += ".txt"
        write_txt(i, o)
    if(v.get()==".md"):
        if (o[-3:])!= ".md": #if the last 3 characters are not the md extension, concat the extension to path
            o += ".md"
        write_md(i, o)
    if(v.get()==".docx"): #if user selected .docx file
        if(o[-5:]) != ".docx": #if the last 5 characters are not the docx extension, concat the extension to path
            o += ".docx"
        write_docx(i,o)
    messagebox.showinfo(title="Extraction Status", message="Note Extraction Complete")

def write_txt(i, o):
    with open(o,"a") as output_file: #create txt file with append mode
        with open(i, "r", encoding='utf-8') as r:
                for line in r:
                    if(line[:13]) == "Course Title:": #course title extraction
                        output_file.write(line[14:] + "\n\n")
                    elif(line[:8] == "Chapter:"):
                        output_file.write(line[9:] + "\n")
                    elif(line[:6]=="Video:"):
                        output_file.write(line[7:] + "\n")
                    elif(re.search("[0-9]{1}:[0-9]{2}:[0-9]{2}", line) != None):
                        output_file.write("\t" + line[19:] + "\n")

def write_md(i,o):
    with open(o,"a") as output_file: #create txt file with append mode
        with open(i, "r", encoding='utf-8') as r:
                for line in r:
                    if(line[:13]) == "Course Title:": #course title extraction
                        output_file.write("# " + line[14:] + "\n\n")
                    elif(line[:8] == "Chapter:"):
                        #chapter_title = line.split(". ", 1)
                        output_file.write("## " + line[9:] + "\n")
                    elif(line[:6]=="Video:"):
                        output_file.write("### " + line[7:] + "\n")
                    elif(re.search("[0-9]{1}:[0-9]{2}:[0-9]{2}", line) != None):
                        output_file.write("- " + line[19:] + "\n")
            
def write_docx(i,o):
    document = Document()
    with open(i, "r", encoding= "utf-8") as r:
        #setting default font for document
        style = document.styles['Normal']
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        #start extraction process
        for line in r:
            if(line[:13]) == "Course Title:":
                p = document.add_paragraph()
                course_title = p.add_run(line[14:])
                course_title.font.size=Pt(14)
                course_title.bold=True
            elif(line[:8] == "Chapter:"):
                chapter_title = document.add_paragraph().add_run(line[9:]).bold=True
            elif(line[:6]=="Video:"):
                p = document.add_paragraph()
                video_title = p.add_run(line[7:])
                video_title.font.color.rgb = RGBColor(74, 119, 198)
                video_title.bold = True
            elif(re.search("[0-9]{1}:[0-9]{2}:[0-9]{2}", line) != None):
                note = document.add_paragraph()
                note.add_run(line[19:]).font.color.rgb = RGBColor(74, 119, 198)
                note.paragraph_format.left_indent = Inches(0.5)

    document.save(o)
import main_UI

    




